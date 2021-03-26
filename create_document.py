'''
@Author: rishi
'''

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm , RGBColor , Inches
from docx.shared import Pt
from tkinter import messagebox

def create(name, age, date, tablet):
    if len(name) == 0:
        messagebox.showinfo("Error", "No patient name is found, please retry again!")
        raise ValueError('No name found')
    # init document
    document = Document()
    para = document.add_heading("")
    # Header
    runn = para.add_run("HealthCare Hospitals \n")
    run = para.add_run(
        " Dr.Rishi Raj M.D.(Neurology) \t \t Dr.Swathy M.S.(Ortho) \t Dr.Shreevarshann M.S.(Ophthalmology) \n No.20,Blossom Avenue,Madura \n Ph:0452-2588 1522 \n")
    # Set page attributes
    font = run.font
    runn.font.size = Pt(18)
    runn.font.color.rgb = RGBColor(153, 17, 150)
    font.size = Pt(16)
    font.color.rgb = RGBColor(217, 17, 213)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    para1 = document.add_paragraph()
    run1 = para1.add_run("Name:{}".format(name))
    run1.font.size = Pt(14)

    para2 = document.add_paragraph()
    run2 = para2.add_run("Age:{}".format(age))
    run2.font.size = Pt(14)

    para3 = document.add_paragraph()
    run3 = para3.add_run("Date:{}".format(date))
    run3.font.size = Pt(14)

    if (len(tablet) == 0):
        messagebox.showinfo("Alert", "Please include atleast one tablet")
    else:
        table = document.add_table(len(tablet), 5)

        heading_cells = table.rows[0].cells
        heading_cells[0].text = 'Name'
        heading_cells[1].text = 'Dosage and Strength'
        heading_cells[2].text = 'Duration and Frequency'

        for i in range(len(tablet)):
            cells = table.add_row().cells
            cells[0].text = tablet[i][0]
            cells[1].text = tablet[i][1]
            cells[2].text = tablet[i][2]

    para5 = document.add_paragraph("\n \n")

    header = document.sections[0].footer
    f1 = header.add_paragraph()

    runf = f1.add_run("\n Signature \n I hereby accept that this prescription was verified")
    runf.font.size = Pt(16)
    runf.font.color.rgb = RGBColor(0, 0, 0)
    f1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    document.save("{}.docx".format(name))
    messagebox.showinfo("Successfully completed",
                        "Success! Please verify the document generated and click upload \n Document Name:{}.docx".format(
                            name))

if __name__ == "__main__":
    create("new_test", "20", "20-10-2020", [["tab1", "10", "3", "After", "1-0-1"]], [])